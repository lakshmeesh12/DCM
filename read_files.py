import re
import docx
import PyPDF2
from pdf2image import convert_from_path
import fitz  # PyMuPDF for robust PDF handling
import utils.logger_test as logger
import pytesseract
from paddleocr import PaddleOCR
import os
import asyncio
from concurrent.futures import ThreadPoolExecutor
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from utils import common
import zipfile
import io
import tempfile
from PIL import Image
import threading
import queue
import uuid

pytesseract.pytesseract.tesseract_cmd = common.pytesserct_path

class ReadFiles:
    """
    This class extracts text from .docx, .pdf, and image files, including text from tables and embedded images,
    with parallel processing for high performance.
    """

    def __init__(self, max_ocr_instances=6):
        self.max_ocr_instances = max_ocr_instances
        self.ocr_pool = queue.Queue()
        self._initialize_ocr_pool()
        self.loop = asyncio.get_event_loop()
        self.executor = ThreadPoolExecutor(max_workers=os.cpu_count() * 2)

    def _initialize_ocr_pool(self):
        """Initialize a pool of PaddleOCR instances for reuse."""
        for _ in range(self.max_ocr_instances):
            try:
                ocr = PaddleOCR(use_angle_cls=True, lang='en', show_log=False)
                self.ocr_pool.put(ocr)
            except Exception as e:
                print(f"Warning: Failed to initialize PaddleOCR instance: {e}")

    def _get_ocr_instance(self):
        """Get a PaddleOCR instance from the pool, or use pytesseract as fallback."""
        try:
            return self.ocr_pool.get_nowait()
        except queue.Empty:
            print("No OCR instances available, falling back to pytesseract.")
            return None

    def _release_ocr_instance(self, ocr):
        """Return an OCR instance to the pool."""
        if ocr is not None:
            self.ocr_pool.put(ocr)

    async def get_text_image(self, data, ocr_instance=None):
        """
        Extract text from an image using OCR asynchronously.

        Parameters:
        ---------
        data: Image file path, bytes, or PIL Image object.
        ocr_instance: PaddleOCR instance (optional).

        Return:
        ------
        text: Extracted text.
        """
        ocr = None
        temp_path = None
        try:
            ocr = ocr_instance or self._get_ocr_instance()
            if ocr is None:
                try:
                    if isinstance(data, bytes):
                        image = Image.open(io.BytesIO(data))
                    elif isinstance(data, Image.Image):
                        image = data
                    else:
                        image = Image.open(data)
                    text = await self.loop.run_in_executor(
                        self.executor, lambda: pytesseract.image_to_string(image, config='--psm 6')
                    )
                    return re.sub(r'\s+', ' ', text).strip()
                except Exception as tesseract_e:
                    print(f"Pytesseract failed: {tesseract_e}")
                    logger.logging.error(f"Pytesseract error: {tesseract_e}")
                    return ""

            # Handle PaddleOCR with unique temporary file
            unique_suffix = f"{uuid.uuid4()}.jpg"
            if isinstance(data, Image.Image):
                image = data.convert('RGB')
                with tempfile.NamedTemporaryFile(suffix=unique_suffix, delete=False) as temp_file:
                    temp_path = temp_file.name
                    image.save(temp_path, 'JPEG')
            elif isinstance(data, bytes):
                image = Image.open(io.BytesIO(data)).convert('RGB')
                with tempfile.NamedTemporaryFile(suffix=unique_suffix, delete=False) as temp_file:
                    temp_path = temp_file.name
                    image.save(temp_path, 'JPEG')
            else:
                temp_path = data  # Assume data is a file path
                if not os.path.exists(temp_path):
                    raise FileNotFoundError(f"Image file not found: {temp_path}")

            try:
                output = await self.loop.run_in_executor(
                    self.executor, lambda: ocr.ocr(temp_path)
                )
                if output and output[0]:
                    texts = [line[1][0] for line in output[0] if line and len(line) >= 2 and line[1] and len(line[1]) >= 1]
                    return " ".join(texts).strip()
                return ""
            finally:
                if temp_path and os.path.exists(temp_path):
                    try:
                        os.unlink(temp_path)
                    except Exception as unlink_e:
                        print(f"Warning: Failed to delete temp file {temp_path}: {unlink_e}")
                        logger.logging.warning(f"Failed to delete temp file {temp_path}: {unlink_e}")

        except Exception as e:
            print(f"Exception in get_text_image: {str(e)}")
            logger.logging.error(f"get_text_image error: {str(e)}")
            return ""
        finally:
            self._release_ocr_instance(ocr)

    async def get_text_docx(self, filename):
        """
        Extract text from a .docx file, including paragraphs, tables, and embedded images asynchronously.

        Parameters:
        ---------
        filename: Input document path.

        Return:
        ------
        allText: Extracted text from paragraphs, tables, and images.
        """
        try:
            print(f"Processing .docx: {filename}")
            if not os.path.exists(filename):
                raise FileNotFoundError(f"File not found: {filename}")

            doc = await self.loop.run_in_executor(self.executor, lambda: docx.Document(filename))
            fullText = []

            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    fullText.append(para.text.strip())
            print(f"Text from paragraphs (first 200 chars): {''.join(fullText)[:200] if fullText else 'No text'}...")

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            fullText.append(cell_text)
            print(f"Text from tables (first 200 chars): {''.join(fullText)[:200] if fullText else 'No text'}...")

            # Extract images from .docx
            async def process_image(file_info, docx_zip):
                with docx_zip.open(file_info) as image_file:
                    image_data = image_file.read()
                    image_text = await self.get_text_image(image_data)
                    return image_text

            try:
                with zipfile.ZipFile(filename) as docx_zip:
                    image_tasks = [
                        process_image(file_info, docx_zip)
                        for file_info in docx_zip.infolist()
                        if file_info.filename.startswith('word/media/')
                    ]
                    image_texts = await asyncio.gather(*image_tasks, return_exceptions=True)
                    for image_text in image_texts:
                        if isinstance(image_text, str) and image_text:
                            fullText.append(image_text)
                            print(f"Text from embedded image: {image_text[:100] if image_text else 'No text'}...")
            except Exception as image_e:
                print(f"Error extracting images from .docx: {image_e}")
                logger.logging.error(f"get_text_docx image extraction error: {str(image_e)}")

            allText = '\n'.join(fullText)
            print(f"Total text from .docx (first 200 chars): {allText[:200] if allText else 'No text'}...")
            return allText

        except Exception as e:
            print(f"Exception in get_text_docx: {str(e)}")
            logger.logging.error(f"get_text_docx error: {str(e)}")
            return ""

    async def extract_embedded_images_from_pdf(self, pdf_document):
        """
        Extract embedded images directly from PDF using PyMuPDF.
        This is much faster than converting pages to images.
        """
        image_tasks = []
        
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            image_list = page.get_images(full=True)
            
            for img_index, img in enumerate(image_list):
                try:
                    # Get the image data
                    xref = img[0]
                    base_image = pdf_document.extract_image(xref)
                    image_bytes = base_image["image"]
                    
                    # Create OCR task for this image
                    task = self.get_text_image(image_bytes)
                    image_tasks.append(task)
                    
                    print(f"Found embedded image {img_index + 1} on page {page_num + 1}")
                    
                except Exception as e:
                    print(f"Error extracting image {img_index + 1} from page {page_num + 1}: {e}")
                    continue
        
        # Process all images in parallel
        if image_tasks:
            image_texts = await asyncio.gather(*image_tasks, return_exceptions=True)
            valid_texts = [text for text in image_texts if isinstance(text, str) and text.strip()]
            return valid_texts
        
        return []

    async def get_text_pdf_page(self, page, page_num):
        """
        Extract text from a single PDF page (native text only).
        Images are handled separately for better performance.
        """
        try:
            # Extract native text from the page
            page_text = await self.loop.run_in_executor(
                self.executor, lambda: page.get_text("text").replace("\n", " ").replace(" -", "-")
            )
            
            if page_text:
                print(f"Page {page_num+1} text length: {len(page_text)}")
                return re.sub(r'\s+', ' ', page_text).strip()
            
            return ""

        except Exception as page_e:
            print(f"Error processing page {page_num+1}: {page_e}")
            logger.logging.error(f"Page {page_num+1} processing error: {str(page_e)}")
            return ""

    async def get_text_pdf(self, filename):
        """
        Extract text from a PDF file, processing pages and images separately for optimal performance.
        """
        try:
            print(f"Processing PDF: {filename}")
            if not os.path.exists(filename):
                raise FileNotFoundError(f"File not found: {filename}")

            with fitz.open(filename) as pdf:
                total_pages = len(pdf)
                print(f"Total pages: {total_pages}")
                
                # Process pages for native text in parallel
                page_tasks = [
                    self.get_text_pdf_page(pdf[page_num], page_num) 
                    for page_num in range(total_pages)
                ]
                
                # Extract embedded images in parallel
                image_extraction_task = self.extract_embedded_images_from_pdf(pdf)
                
                # Wait for both text and image extraction to complete
                page_texts, image_texts = await asyncio.gather(
                    asyncio.gather(*page_tasks, return_exceptions=True),
                    image_extraction_task,
                    return_exceptions=True
                )
                
                # Combine all text
                all_text = []
                
                # Add page texts
                if isinstance(page_texts, list):
                    all_text.extend([t for t in page_texts if isinstance(t, str) and t])
                
                # Add image texts
                if isinstance(image_texts, list):
                    all_text.extend(image_texts)
                    print(f"Extracted text from {len(image_texts)} embedded images")
                
                final_text = ' '.join(all_text)
                print(f"Full extracted text from PDF ({len(final_text)} chars)")
                return re.sub(r'\s+', ' ', final_text).strip()

        except Exception as e:
            print(f"Exception in get_text_pdf: {str(e)}")
            logger.logging.error(f"get_text_pdf error: {str(e)}")
            return ""

    async def process_file(self, file_path):
        """
        Process a single file asynchronously.

        Parameters:
        ---------
        file_path: Input file path.

        Return:
        ------
        tuple: (file_path, extracted_text)
        """
        try:
            if not file_path or not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found or invalid: {file_path}")

            file_extension = file_path.split(".")[-1].lower()
            print(f"Processing file type: {file_extension} for {file_path}")

            if file_extension == "txt":
                encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
                for encoding in encodings:
                    try:
                        with open(file_path, 'r', encoding=encoding) as file_to_read:
                            text = await self.loop.run_in_executor(
                                self.executor, lambda: file_to_read.read()
                            )
                            text = re.sub(r"\s+", " ", text).strip()
                            print(f"Full extracted text from txt ({len(text)} chars):\n{text}")
                            return file_path, text
                    except UnicodeDecodeError:
                        continue
                with open(file_path, 'rb') as file_to_read:
                    raw_data = await self.loop.run_in_executor(
                        self.executor, lambda: file_to_read.read()
                    )
                    text = raw_data.decode('utf-8', errors='ignore')
                    text = re.sub(r"\s+", " ", text).strip()
                    return file_path, text

            elif file_extension in ["docx", "doc"]:
                text = await self.get_text_docx(file_path)
                return file_path, text

            elif file_extension == "pdf":
                text = await self.get_text_pdf(file_path)
                return file_path, text

            elif file_extension in ["jpeg", "jpg", "png", "jp2", "pbm", "ppm", "tif", "tiff", "bmp", "gif"]:
                text = await self.get_text_image(file_path)
                print(f"Full extracted text from image ({len(text)} chars):\n{text}")
                return file_path, text

            else:
                error_msg = f"Unsupported file format: {file_extension}"
                print(error_msg)
                logger.logging.error(error_msg)
                return file_path, ""

        except Exception as e:
            error_msg = f"file_reader error for {file_path}: {str(e)}"
            print(error_msg)
            logger.logging.error(error_msg)
            return file_path, ""

    async def file_reader(self, filename):
        """
        Extract text from a file or list of files based on extension, processing in parallel.

        Parameters:
        ---------
        filename: Input file path or list of file paths.

        Return:
        ------
        results: Dictionary mapping filenames to extracted text or single text string.
        """
        print(f"Ready to read: {filename}")
        try:
            filenames = [filename] if isinstance(filename, str) else filename
            results = {}

            tasks = [self.process_file(f) for f in filenames]
            file_results = await asyncio.gather(*tasks, return_exceptions=True)
            for file_path, text in file_results:
                if isinstance(text, str):
                    results[os.path.basename(file_path)] = text

            return results[os.path.basename(filename)] if isinstance(filename, str) else results

        except Exception as e:
            error_msg = f"file_reader error: {str(e)}"
            print(error_msg)
            logger.logging.error(error_msg)
            return "" if isinstance(filename, str) else {}

    def __del__(self):
        """Clean up resources."""
        self.executor.shutdown(wait=True)