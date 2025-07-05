import glob
import shutil
import os
import pypdf
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image, ImageDraw, ImageFont
import re

temp_folder = "temp_files_to_be_extracted"
global_entities = [
    'CREDIT_CARD', 'CRYPTO', 'DATE_TIME', 'EMAIL_ADDRESS', 'IBAN_CODE', 'IP_ADDRESS', 'NRP',
    'LOCATION', 'PERSON', 'PHONE_NUMBER', 'MEDICAL_LICENSE', 'URL'
]

custom_entities = [
    "IN_PHONE_NUMBER", "IN_CREDIT_CARD", "IN_AADHAR_CARD_CUSTOM", "IN_PASSPORT_CUSTOM",
    "IN_VEHICLE_REGISTRATION_CUSTOM", "IN_VOTER_ID_CUSTOM", "IBAN_CODE_CUSTOM", "CRYPTO_CUSTOM",
    "MEDICAL_LICENSE_CUSTOM", "IN_PAN_CUSTOM", "IN_GST_NUMBER", "IN_UPI_ID",
    "IN_BANK_ACCOUNT", "IN_IFSC_CODE", "IN_DRIVING_LICENSE"
]

country_entities = {
    'USA': ['US_BANK_NUMBER', 'US_DRIVER_LICENSE', 'US_ITIN', 'US_PASSPORT', 'US_SSN'],
    'UK': ['UK_NHS', 'UK_NINO'],
    'Spain': ['ES_NIF', 'ES_NIE'],
    'Italy': ['IT_FISCAL_CODE', 'IT_DRIVER_LICENSE', 'IT_VAT_CODE', 'IT_PASSPORT', 'IT_IDENTITY_CARD'],
    'Poland': ['PL_PESEL'],
    'Singapore': ['SG_NRIC_FIN', 'SG_UEN'],
    'Australia': ['AU_ABN', 'AU_ACN', 'AU_TFN', 'AU_MEDICARE'],
    'India': [
        'IN_PAN', 'IN_AADHAR', 'IN_VEHICLE_REGISTRATION', 'IN_VOTER', 'IN_PASSPORT',
        'IN_PHONE_NUMBER', 'IN_CREDIT_CARD', 'IN_AADHAR_CARD_CUSTOM', 'IN_PASSPORT_CUSTOM',
        'IN_VEHICLE_REGISTRATION_CUSTOM', 'IN_VOTER_ID_CUSTOM', 'IN_PAN_CUSTOM',
        'IN_GST_NUMBER', 'IN_UPI_ID', 'IN_BANK_ACCOUNT', 'IN_IFSC_CODE', 'IN_DRIVING_LICENSE'
    ],
    'Finland': ['FI_PERSONAL_IDENTITY_CODE']
}

allowed_extensions = {'txt', 'pdf', 'png', 'jpg', 'jpeg', 'docx', 'tif', 'doc'}
dataframe_columns = ['File Name', 'text']
poppler_path = r"C:\poppler-23.01.0\Library\bin"
pytesserct_path = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# Updated and improved regex patterns
regex_for_in_phone_number = r'(?:\+91[\s\-]?)?(?:0)?[6-9]\d{9}\b'
regex_for_credit_card = r'(?<!\d)(?:4[0-9]{3}[ \-]?[0-9]{4}[ \-]?[0-9]{4}[ \-]?[0-9]{4}|5[1-5][0-9]{2}[ \-]?[0-9]{4}[ \-]?[0-9]{4}[ \-]?[0-9]{4}|3[47][0-9]{2}[ \-]?[0-9]{6}[ \-]?[0-9]{5}|3(?:0[0-5]|[68][0-9])[ \-]?[0-9]{6}[ \-]?[0-9]{4}|6(?:011|5[0-9]{2})[ \-]?[0-9]{4}[ \-]?[0-9]{4}[ \-]?[0-9]{4}|(?:2131|1800|35\d{3})[ \-]?\d{11})(?!\d)'

# Enhanced Aadhar patterns
regex_for_custom_in_aadhar = r'\b(?:\d{4}[\s\-]?\d{4}[\s\-]?\d{4}|\d{12})\b'

# Enhanced PAN pattern
regex_for_in_pan = r'\b[A-Z]{5}\d{4}[A-Z]\b'

# Enhanced Voter ID pattern
regex_for_custom_in_voter = r'\b[A-Z]{3}\d{7}\b'

# Enhanced Bank Account pattern
regex_for_in_bank_account = r'(?<!\d)\d{12,18}(?!\d)'

# Enhanced Phone pattern
regex_for_in_phone_number = r'(?:\+91[\s\-]?)?(?:0)?[6-9]\d{9}\b'

# Enhanced Passport pattern
regex_for_custom_in_passport = r'\b[A-PR-WY][0-9]{7}\b'

# Enhanced Driving License pattern
regex_for_in_driving_license = r'\b[A-Z]{2}-?\d{2}-?\d{4}-?\d{7}\b'

# Keep other existing patterns as they are
regex_for_custom_in_vehicle_registration = r'\b[A-Z]{2}[ -]?\d{1,2}[ -]?[A-Z]{1,3}[ -]?\d{4}\b'
regex_for_custom_in_iban_code = r'\b[A-Z]{2}[0-9]{2}(?:[ \-]?[A-Z0-9]{4}){2,7}(?:[ \-]?[A-Z0-9]{1,3})?\b'
regex_for_custom_crypto = r'\b(?:bc1|[13])[a-zA-HJ-NP-Z0-9]{25,39}\b|\b0x[a-fA-F0-9]{40}\b'
regex_for_custom_medical_license = r'\b[A-Z]{2,5}/\d{1,6}\b'
regex_for_in_gst_number = r'\b[0-9]{2}[A-Z]{5}[0-9]{4}[A-Z]{1}[1-9A-Z]{1}[0-9A-Z]{1}[0-9A-Z]{1}\b'
regex_for_in_upi_id = r'[\w.\-_]{2,256}@[a-zA-Z]{2,64}\b'
regex_for_in_ifsc_code = r'\b[A-Z]{4}0[A-Z0-9]{6}\b'


# Directory setup
base_dir = os.path.dirname(os.path.abspath(__file__))
watermarked_output_folder = os.path.join(base_dir, "..", "..", "watermarked_files")
os.makedirs(watermarked_output_folder, exist_ok=True)

header_output_folder = os.path.join(base_dir, "..", "..", "header_output_files")
os.makedirs(header_output_folder, exist_ok=True)

entities_list = global_entities + [entity for entities in country_entities.values() for entity in entities] + custom_entities

def delete_folder_files(path, new_dir: bool, formats: tuple = None):
    """
    Deletes files or folder. If formats are passed, only files of those formats are deleted.

    Parameters:
    -----------
    path (str): Path to the directory or file
    new_dir (bool): Creates a new folder based on the path
    formats (tuple): File formats to be deleted

    Returns:
    --------
    bool: True if successfully executed, else False
    """
    try:
        if os.path.isfile(path) or os.path.islink(path):
            os.remove(path)
            print("file removed - ", path)
        elif os.path.isdir(path) and formats is None:
            shutil.rmtree(path)
            print('Directory Deleted - ', path)
            if new_dir:
                os.mkdir(path)
                print('New Directory Created - ', path)
        elif os.path.isdir(path) and formats is not None:
            files = glob.glob(os.path.join(path, "*"))
            for file in files:
                if file.endswith(formats):
                    print("deleting file - ", file)
                    os.remove(file)
        elif not os.path.isdir(path) and new_dir:
            os.mkdir(path)
            print('New Directory Created - ', path)
        return True
    except Exception as e:
        print('Exception occurred while deleting file/folder - ', str(e))
        return False

def add_watermark_pdf(input_file_name, watermark_text, temp_folder="temp_files_to_be_extracted", isFrom_docx=False):
    base_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "../../"))
    input_pdf_path = os.path.join(base_dir, temp_folder, input_file_name)

    existing_pdf = pypdf.PdfReader(input_pdf_path)
    output = pypdf.PdfWriter()

    for page_num in range(len(existing_pdf.pages)):
        page = existing_pdf.pages[page_num]
        page_width = float(page.mediabox.width)
        page_height = float(page.mediabox.height)

        packet = io.BytesIO()
        can = canvas.Canvas(packet, pagesize=(page_width, page_height))
        can.setFont("Helvetica-Bold", 80)
        can.setFillColorRGB(0.5, 0.5, 0.5)
        can.setFillAlpha(0.3)
        can.saveState()
        can.translate(page_width / 2, page_height / 2)
        can.rotate(45)
        can.drawCentredString(0, 0, watermark_text)
        can.restoreState()
        can.save()

        packet.seek(0)
        watermark_pdf = pypdf.PdfReader(packet)
        page.merge_page(watermark_pdf.pages[0])
        output.add_page(page)

    output_pdf_path = os.path.join(base_dir, "watermarked_files", input_file_name)
    with open(output_pdf_path, "wb") as output_stream:
        output.write(output_stream)
        print(f"Watermark output path is: {output_pdf_path}")
    
    return output_pdf_path

def add_watermark_docx(input_file_name, watermark_text, convert_back_to_docx=False):
    # Define paths
    base_dir = os.path.dirname(__file__)
    input_folder = os.path.join(base_dir, '..', '..', 'temp_files_to_be_extracted')
    # watermarked_output_folder = os.path.join(base_dir, 'watermarked_output')
    temp_folder = os.path.join(base_dir, 'temp_pdf')
   
    # os.makedirs(watermarked_output_folder, exist_ok=True)
    os.makedirs(temp_folder, exist_ok=True)
   
    # Input and temporary PDF paths
    input_docx_path = os.path.join(input_folder, input_file_name)
    temp_pdf_name = input_file_name.replace('.docx', '.pdf')
    temp_pdf_path = os.path.join(temp_folder, temp_pdf_name)
   
    # Convert DOCX to PDF
    print(f"Converting {input_docx_path} to PDF...")
    docx_to_pdf(input_docx_path, temp_pdf_path)
   
    # Apply watermark to PDF
    print(f"Applying watermark to {temp_pdf_path}...")
    watermarked_pdf_path = add_watermark_pdf(temp_pdf_name, watermark_text, temp_folder = temp_folder, isFrom_docx = True)
    print("Returned Output Path")
 
    # Optionally convert back to DOCX
    if convert_back_to_docx:
        output_docx_path = os.path.join(watermarked_output_folder, input_file_name)
        print(f"Converting watermarked PDF back to DOCX: {output_docx_path}")
        cv = Converter(watermarked_pdf_path)
        cv.convert(output_docx_path)
        cv.close()
        print(f"Watermarked DOCX saved at: {output_docx_path}")
    else:
        print(f"Watermarked PDF saved at: {watermarked_pdf_path}")
 
 

def add_header_pdf(input_file_name, header_text, text_color='#000000', bg_color='#FFFF00'):
    base_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "../../"))
    input_pdf_path = os.path.join(base_dir, "temp_files_to_be_extracted", input_file_name)
    file_extension = input_file_name.split(".")[-1].lower()
    if file_extension == 'txt':
        return input_pdf_path

    existing_pdf = pypdf.PdfReader(input_pdf_path)
    output = pypdf.PdfWriter()
    text_rgb = hex_to_rgb(text_color)
    bg_rgb = hex_to_rgb(bg_color)

    for page_num in range(len(existing_pdf.pages)):
        page = existing_pdf.pages[page_num]
        page_width = float(page.mediabox.width)
        page_height = float(page.mediabox.height)

        packet = io.BytesIO()
        can = canvas.Canvas(packet, pagesize=(page_width, page_height))
        can.setFont("Helvetica-Bold", 12)
        text_width = can.stringWidth(header_text, "Helvetica-Bold", 12)
        margin_right = 20
        margin_top = 20
        x_position = page_width - margin_right - text_width
        y_position = page_height - margin_top
        padding = 2

        can.setFillColorRGB(bg_rgb[0]/255, bg_rgb[1]/255, bg_rgb[2]/255)
        can.rect(
            x_position - padding,
            y_position - padding,
            text_width + 2 * padding,
            12 + 2 * padding,
            fill=1
        )
        can.setStrokeColorRGB(bg_rgb[0]/255, bg_rgb[1]/255, bg_rgb[2]/255)
        can.setLineWidth(1)
        can.rect(
            x_position - padding,
            y_position - padding,
            text_width + 2 * padding,
            12 + 2 * padding,
            fill=0
        )
        can.setFillColorRGB(text_rgb[0]/255, text_rgb[1]/255, text_rgb[2]/255)
        can.drawString(x_position, y_position, header_text)
        can.save()

        packet.seek(0)
        header_pdf = pypdf.PdfReader(packet)
        page.merge_page(header_pdf.pages[0])
        output.add_page(page)

    output_pdf_path = os.path.join(base_dir, "header_output_files", input_file_name)
    with open(output_pdf_path, "wb") as output_stream:
        output.write(output_stream)
        print(f"Header output path is: {output_pdf_path}")

    return output_pdf_path

def add_header_docx(input_file_name, header_text, text_color='#000000', bg_color='#FFFF00'):
    base_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "../../"))
    input_docx_path = os.path.join(base_dir, "temp_files_to_be_extracted", input_file_name)

    doc = Document(input_docx_path)
    text_rgb = hex_to_rgb(text_color)
    bg_rgb = hex_to_rgb(bg_color)

    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            paragraph.clear()
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = header_para._element
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), bg_color.lstrip('#'))
        p.get_or_add_pPr().append(shd)
        run = header_para.add_run(header_text)
        run.font.name = 'Helvetica'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(*text_rgb)
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        for border in ['w:top', 'w:right', 'w:bottom', 'w:left']:
            bdr = OxmlElement(border)
            bdr.set(qn('w:val'), 'single')
            bdr.set(qn('w:sz'), '2')
            bdr.set(qn('w:color'), bg_color.lstrip('#'))
            pBdr.append(bdr)
        pPr.append(pBdr)
        section.header_distance = Inches(0.28)
        section.right_margin = Inches(0.28)

    output_docx_path = os.path.join(base_dir, "header_output_files", input_file_name)
    doc.save(output_docx_path)
    print(f"Header output path is: {output_docx_path}")
    return output_docx_path

def add_header_image(input_file_name, header_text, text_color='#000000', bg_color='#FFFF00'):
    base_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "../../"))
    input_image_path = os.path.join(base_dir, "temp_files_to_be_extracted", input_file_name)

    image = Image.open(input_image_path)
    image = image.convert("RGB")
    draw = ImageDraw.Draw(image)
    image_width, image_height = image.size

    try:
        font = ImageFont.truetype("arialbd.ttf", 20)
    except IOError:
        font = ImageFont.load_default()
        print("Warning: Helvetica-Bold not found, using default font")

    text_bbox = draw.textbbox((0, 0), header_text, font=font)
    text_width = text_bbox[2] - text_bbox[0]
    text_height = text_bbox[3] - text_bbox[1]
    margin_right = 20
    margin_top = 20
    x_position = image_width - text_width - margin_right
    y_position = margin_top
    text_rgb = hex_to_rgb(text_color)
    bg_rgb = hex_to_rgb(bg_color)
    padding = 2

    draw.rectangle(
        (
            x_position - padding,
            y_position - padding,
            x_position + text_width + padding,
            y_position + text_height + padding
        ),
        fill=bg_rgb
    )
    draw.rectangle(
        (
            x_position - padding,
            y_position - padding,
            x_position + text_width + padding,
            y_position + text_height + padding
        ),
        outline=bg_rgb,
        width=1
    )
    draw.text((x_position, y_position), header_text, font=font, fill=text_rgb)

    output_image_path = os.path.join(base_dir, "header_output_files", input_file_name)
    image.save(output_image_path)
    print(f"Header output path is: {output_image_path}")
    return output_image_path

def hex_to_rgb(hex_color):
    """
    Convert hex color (e.g., '#12af42') to RGB tuple (e.g., (18, 175, 66)).
    Raises ValueError for invalid hex color.
    """
    hex_color = hex_color.lstrip('#')
    if not re.match(r'^[0-9a-fA-F]{6}$', hex_color):
        raise ValueError(f"Invalid hex color: #{hex_color}")
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))