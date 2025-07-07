import concurrent.futures
import json
import os
from fastapi import FastAPI, Request, Response, HTTPException, File, UploadFile, Form, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse
from werkzeug.utils import secure_filename
import pandas as pd
from TemporaryStoringClasses import SetGetSelectedEntities, SetGetAWSDetails, SetGetDataFrame, SetGetAzureDetails, StreamableFiles, SetGetStreamedDataFrame, CloudFiles, QAData, ChatHistory, VectorDB
from parse_pii import ParsePii
from read_files import ReadFiles
from table_extraction.table_extraction_main import main as te_main
from utils import common
from io import BytesIO
from zipfile import ZipFile
import glob
import requests
import boto3
import PyPDF2
import docx
from azure.storage.blob import BlobServiceClient
from dotenv import find_dotenv, load_dotenv
import re
import traceback
from datetime import datetime
from utils.common import add_watermark_pdf, add_watermark_docx, add_header_pdf, add_header_docx, add_header_image
from pydantic import BaseModel
from typing import List
from TemporaryStoringClasses import NeededInformation
from parse_pii import needed_information
from google_drive import GoogleDriveClient
from llm import LLMEntityDetector
from concurrent.futures import ThreadPoolExecutor

# Initialize FastAPI app
app = FastAPI()

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS"],
    allow_headers=["Content-Type", "Authorization"],
    expose_headers=["Content-Disposition"]
)

# Existing configurations
currFolder = common.temp_folder
ALLOWED_EXTENSIONS = common.allowed_extensions

# Existing storage objects
selected = SetGetSelectedEntities()
setGetBucketName = SetGetAWSDetails()
setGetAzureDetails = SetGetAzureDetails()
setGetDataframe = SetGetDataFrame()
setGetStreamingDataframe = SetGetStreamedDataFrame()
streamable_files = StreamableFiles()
temp_cloud_files = CloudFiles()
qadata = QAData()
chat_history_obj = ChatHistory()
vect_db_obj = VectorDB()

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

async def copy_file(file: UploadFile, filename: str):
    file_location = os.path.join("..", currFolder, filename)
    contents = await file.read()
    with open(file_location, "wb") as f:
        f.write(contents)

async def copy_files(files):
    saved_files = []
    for file in files:
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            try:
                await copy_file(file, filename)
                saved_files.append(filename)
                print(f"Saved file: {filename}")
            except Exception as e:
                print(f"Failed to save file {filename}: {str(e)}")
        else:
            print(f"Skipped file {file.filename}: Invalid extension or empty")
    return saved_files

@app.get('/')
async def user_login():
    return JSONResponse({"message": "Login endpoint. Please use the React frontend to access."})

@app.post('/upload')
async def upload_form(request: Request):
    if request.method == 'POST':
        empty_df = pd.DataFrame()
        setGetStreamingDataframe.set_stream_dataframe(empty_df)
        setGetDataframe.set_dataframe(empty_df)
        common.delete_folder_files(currFolder, True)
        return JSONResponse({"message": "Upload form initialized", "status": "success"})
    return JSONResponse({"message": "GET request not supported for upload form"})

@app.post("/file_handler")
async def fileHandler(files: List[UploadFile] = File(...)):
    print("Request files:", [f.filename for f in files])
    try:
        if not files:
            print("No files found in request")
            raise HTTPException(
                status_code=400,
                detail={
                    "status": "error",
                    "message": "No files found in request. Expected key: 'files' or 'files[]'",
                    "visibility": "hidden"
                }
            )

        valid_files = [f for f in files if f.filename and allowed_file(f.filename)]
        print(f"Valid files: {[f.filename for f in valid_files]}")
        print(f"ALLOWED_EXTENSIONS: {ALLOWED_EXTENSIONS}")

        if not valid_files:
            print("No valid files after filtering")
            raise HTTPException(
                status_code=400,
                detail={
                    "status": "error",
                    "message": f"No valid files uploaded. Ensure files have supported extensions: {', '.join(ALLOWED_EXTENSIONS)}",
                    "visibility": "hidden"
                }
            )
        currFolder_renamed = os.path.join("..", currFolder)
        common.delete_folder_files(currFolder_renamed, True)
        saved_filenames = await copy_files(valid_files)

        if saved_filenames:
            return JSONResponse({
                "status": "success",
                "message": f"{len(saved_filenames)} files successfully uploaded",
                "visibility": "visible",
                "filenames": saved_filenames
            })
        else:
            print("No files were saved")
            raise HTTPException(
                status_code=500,
                detail={
                    "status": "error",
                    "message": "Failed to save any files. Check file formats and permissions.",
                    "visibility": "hidden"
                }
            )

    except Exception as e:
        print(f"Error in file_handler: {str(e)}")
        print(traceback.format_exc())
        raise HTTPException(
            status_code=500,
            detail={
                "status": "error",
                "message": f"Server error: {str(e)}",
                "visibility": "hidden"
            }
        )

def extractedDataFrame(entities, user_option):
    try:
        output_path = os.path.join(os.getcwd(), "images")
        common.delete_folder_files(output_path, True)
        parse_pii = ParsePii()
        read_files = ReadFiles()
        dataFrame = pd.DataFrame(columns=common.dataframe_columns)
        current_dir = os.path.join("..", common.temp_folder)
        files = os.listdir(current_dir)
        paths = [os.path.join(current_dir, f) for f in files]
        print(f"Paths to process: {paths}")

        if user_option == "TablesExtraction":
            result = te_main(paths, output_path)
            return {
                "status": "success",
                "body": result
            }

        dataFrame["File Name"] = paths
        dataFrame["text"] = dataFrame["File Name"].apply(lambda path: read_files.file_reader(path))
        dataFrame["File Name"] = dataFrame["File Name"].apply(lambda path: os.path.basename(path))
        stream_df = setGetStreamingDataframe.get_stream_dataframe()
        if not stream_df.empty:
            dataFrame = pd.concat([dataFrame, stream_df], ignore_index=True)
        dataFrame["Has PII ?"] = dataFrame["text"].apply(lambda text: parse_pii.classify(text, entities=entities))
        dataFrame["Extracted Fields"] = dataFrame["text"].apply(lambda text: parse_pii.extract(text, entities=entities))
        dataFrame.fillna("", inplace=True)
        dataFrame = dataFrame[dataFrame['text'] != ""]

        # Handle None return from parse_pii.extract
        def safe_extract(text, entities):
            result = parse_pii.extract(text, entities=entities)
            return result[1] if result is not None and isinstance(result, (list, tuple)) and len(result) > 1 else {}

        dataFrame["Needed Information"] = dataFrame["text"].apply(lambda text: safe_extract(text, entities))
        dataFrame.reset_index(inplace=True, drop=True)
        setGetDataframe.set_dataframe(dataFrame)
        
        return {
            "status": "success",
            "data": {
                "extracted": dataFrame[['File Name', 'Has PII ?', "Extracted Fields"]].to_dict(orient='records'),
                "needed": dataFrame[['File Name', 'Needed Information']].to_dict(orient='records')
            }
        }
    except Exception as e:
        print(traceback.format_exc())
        return {"status": "error", "message": str(e)}

@app.get("/files/{file_name}")
async def download_file(file_name: str):
    base_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    file_extension = file_name.split(".")[-1].lower()
    print(base_dir)
    if file_extension == 'txt':
        file_path = os.path.join(base_dir, "temp_files_to_be_extracted", file_name)
        media_type = 'application/pdf'
    elif file_extension in ['xlsx', 'xls']:
        file_path = os.path.join(base_dir, "entity_extraction", "images", file_name)
        media_type = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    else:
        file_path = os.path.join(base_dir, "header_output_files", file_name)
        media_type = 'application/pdf'
    print("Input Folder...:", base_dir, "-----", file_path)

    print(f"Attempting to serve file: {file_path}")
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        return JSONResponse(status_code=404, content={"error": "File not found"})
    return FileResponse(path=file_path, filename=file_name, media_type=media_type)

@app.post('/mark_document')
def mark_files(file_name: str = Form(...)):
    print("In Marking File...")
    try:
        # Retrieve final_json from NeededInformation
        final_json = needed_information.get_needed_information()
        if not final_json:
            raise HTTPException(
                status_code=500,
                detail={
                    "status": "error",
                    "message": "No PII data available. Please run entity extraction first."
                }
            )

        file_extension = file_name.split(".")[-1].lower()
        for entity in final_json:
            if entity['File Name'] == file_name:
                watermark_text = ""
                print(entity.keys(), entity["Categories"], entity["File Name"])
                if len(entity["Categories"]["Confidential"]) > 2:
                    watermark_text = "Confidential"
                elif len(entity["Categories"]["Private"]) > 2:
                    watermark_text = "Private"
                elif len(entity["Categories"]["Restricted"]):
                    watermark_text = "Restricted"
                else:
                    watermark_text = "Public"

                output_watermark_file_path = ""
                output_header_file_path = ""
                if file_extension == 'pdf':
                    print("File Extension is PDF...")
                    output_watermark_file_path = add_watermark_pdf(file_name, watermark_text=watermark_text)
                    output_header_file_path = add_header_pdf(file_name, header_text=watermark_text)
                elif file_extension in ['docx', 'doc']:
                    output_header_file_path = add_header_docx(file_name, watermark_text)
                elif file_extension in ['png', 'jpg', 'jpeg', 'bmp', 'tiff', 'gif']:
                    output_header_file_path = add_header_image(file_name, watermark_text)
                elif file_extension == 'txt':
                    output_header_file_path = add_header_pdf(file_name, header_text=watermark_text)
                    print("Can't Add Header Mark...")
                else:
                    print(f"Unsupported file type: {file_extension}")

                return {
                    "message": f"{file_name} is marked as {watermark_text}",
                    "header_path": output_header_file_path,
                    "watermark_path": output_watermark_file_path,
                }
        
        raise HTTPException(
            status_code=404,
            detail={
                "status": "error",
                "message": f"File {file_name} not found in PII data."
            }
        )
    except Exception as e:
        print(f"Error in /mark_document: {traceback.format_exc()}")
        raise HTTPException(
            status_code=500,
            detail={
                "status": "error",
                "message": str(e)
            }
        )
def map_ui_entities_to_detector_entities(ui_entities):
    """Map UI entity names to actual detector entity names"""
    entity_mapping = {
        'IN_AADHAR': ['IN_AADHAR_CARD_CUSTOM', 'IN_AADHAAR'],
        'IN_VOTER': ['IN_VOTER_ID_CUSTOM'],
        'IN_PAN': ['IN_PAN_CUSTOM', 'IN_PAN'],
        'IN_PASSPORT': ['IN_PASSPORT_CUSTOM'],
        'IN_VEHICLE_REGISTRATION': ['IN_VEHICLE_REGISTRATION_CUSTOM'],
        'IBAN_CODE': ['IBAN_CODE_CUSTOM', 'IBAN_CODE'],
        'CRYPTO': ['CRYPTO_CUSTOM', 'CRYPTO'],
        'MEDICAL_LICENSE': ['MEDICAL_LICENSE_CUSTOM', 'MEDICAL_LICENSE'],
        'IN_BANK_ACCOUNT': ['IN_BANK_ACCOUNT'],
        'IN_PHONE_NUMBER': ['IN_PHONE_NUMBER', 'PHONE_NUMBER'],
        'IN_CREDIT_CARD': ['IN_CREDIT_CARD', 'CREDIT_CARD'],
        'IN_GST_NUMBER': ['IN_GST_NUMBER'],
        'IN_UPI_ID': ['IN_UPI_ID'],
        'IN_IFSC_CODE': ['IN_IFSC_CODE'],
        'IN_DRIVING_LICENSE': ['IN_DRIVING_LICENSE']
    }
    
    mapped_entities = []
    for entity in ui_entities:
        if entity in entity_mapping:
            mapped_entities.extend(entity_mapping[entity])
        else:
            mapped_entities.append(entity)
    
    return list(set(mapped_entities))  # Remove duplicates


@app.post("/success")
async def displayExtractedResults(
    selectedOption: str = Form(...),
    country: str = Form(None),
    multiple: list[str] = Form(default=[]),
    categoryMapping: str = Form(None),
    user_prompt: str = Form(None)
):
    try:
        # Log raw form data for debugging
        print(f"Received form data: selectedOption={selectedOption}, country={country}, multiple={multiple}, categoryMapping={categoryMapping}, user_prompt={user_prompt}")

        # Initialize ReadFiles and LLMEntityDetector
        read_files = ReadFiles()
        llm_detector = LLMEntityDetector()

        # Parse category mapping if provided
        category_mapping = {}
        if categoryMapping:
            try:
                category_mapping = json.loads(categoryMapping)
                if not isinstance(category_mapping, dict):
                    raise ValueError("categoryMapping must be a JSON object")
            except json.JSONDecodeError:
                raise HTTPException(
                    status_code=400,
                    detail={
                        "status": "error",
                        "message": "Invalid categoryMapping: must be a valid JSON object"
                    }
                )

        # List files in temporary directory
        current_dir = os.path.join("..", common.temp_folder)
        files = [f for f in os.listdir(current_dir) if not f.startswith('~$')]
        paths = [os.path.join(current_dir, f) for f in files]
        print(f"Paths to process: {paths}")

        # Determine entities to detect
        entities = multiple if multiple else []
        print(f"Initial entities from multiple: {entities}")

        if not entities and country:
            country = country.capitalize()
            if country in common.country_entities:
                entities = common.global_entities + common.country_entities[country]
            else:
                raise HTTPException(
                    status_code=400,
                    detail={
                        "status": "error",
                        "message": f"Invalid country: {country}. Supported countries: {list(common.country_entities.keys())}"
                    }
                )
        elif not entities:
            entities = common.entities_list

        print(f"Entities to detect: {entities}")

        # Validate entities
        if not isinstance(entities, list) or entities is None:
            print(f"Error: entities is invalid: {entities}")
            entities = []
            print("Fallback to empty entities list")

        # Process files in parallel
        with ThreadPoolExecutor() as executor:
            results = list(executor.map(
                lambda path: (os.path.basename(path), read_files.file_reader(path)),
                paths
            ))

        # Create DataFrame from results
        dataFrame = pd.DataFrame(results, columns=["File Name", "text"])
        
        # Log extracted text for debugging
        for _, row in dataFrame.iterrows():
            print(f"Extracted text from {row['File Name']} (first 200 chars): {row['text'][:200]}")
            if len(row['text']) > 200:
                print(f"Full extracted text length: {len(row['text'])} chars")

        # Filter out empty text rows
        dataFrame = dataFrame[dataFrame["text"] != ""]
        dataFrame.reset_index(inplace=True, drop=True)

        if dataFrame.empty:
            print("No valid files found with text content: ", paths)
            return {
                "status": "success",
                "data": []
            }

        # Detect and categorize entities in extracted text
        dataFrame["entities"] = dataFrame.apply(
            lambda row: llm_detector.detect_entities(row["text"], entities, category_mapping, row["File Name"], user_prompt),
            axis=1
        )

        # Prepare response: filter empty entities and add has_pii
        response_data = []
        for _, row in dataFrame.iterrows():
            file_data = {
                "File Name": row["File Name"],
                "entities": {
                    category: {
                        entity: values for entity, values in entities.items() if values
                    } for category, entities in row["entities"].items() if any(values for values in entities.values())
                }
            }
            file_data["has_pii"] = "yes" if any(
                values for category in file_data["entities"].values() for values in category.values()
            ) else "no"
            response_data.append(file_data)

        return {
            "status": "success",
            "data": response_data
        }

    except Exception as e:
        print(f"Error in /success: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail={
                "status": "error",
                "message": str(e)
            }
        )

@app.post("/getting_files_from_aws_s3")
async def GettingFilesFromAWSS3(
    bucket: str = Form(...),
    accessKey: str = Form(...),
    secretKey: str = Form(...)
):
    try:
        aws_details = {'bucket': bucket, 'accessKey': accessKey, 'secretKey': secretKey}
        setGetBucketName.set_bucket_name(aws_details['bucket'])
        setGetBucketName.set_access_key(aws_details['accessKey'])
        setGetBucketName.set_secret_key(aws_details['secretKey'])
        url = f"https://w463ug1m4g.execute-api.us-east-1.amazonaws.com/prod/file_extraction/source/{aws_details['bucket']}"
        headers = {"id": aws_details['accessKey'], "key": aws_details['secretKey']}
        files_list_response = requests.get(url, headers=headers)
        if files_list_response.status_code != 200:
            raise HTTPException(
                status_code=400,
                detail={"status": "error", "message": "Please enter valid details"}
            )
        files_list = json.loads(files_list_response.content.decode())
        files = [file[0] for file in files_list]
        smaller_files = [file[0] for file in files_list if file[1] <= 150 and file[0].split('.')[-1] in ['txt', 'pdf', 'docx']]
        temp_cloud_files.set_files(files)
        streamable_files.add_streamable_files(smaller_files)
        return JSONResponse({"status": "success", "files": files})
    except Exception as e:
        raise HTTPException(
            status_code=400,
            detail={"status": "error", "message": str(e)}
        )


@app.post("/aws_selected_files")
async def AwsSelectedFiles(files: List[str] = Form(...)):
    try:
        # ======================================
        currFolder = common.temp_folder
        print("currFolder", currFolder)
        common.delete_folder_files(os.path.join("..",currFolder), True)
        # ======================================

        selected_files_list = files
        if not selected_files_list:
            raise HTTPException(
                status_code=400,
                detail={
                    "status": "error",
                    "message": "Please select a file and click upload!",
                    "files": temp_cloud_files.get_files()
                }
            )
        aws_details = setGetBucketName.get_()
        small_files = streamable_files.get_streamable_files()
        large_files = []
        dataFrame = pd.DataFrame(columns=common.dataframe_columns)
        url = f"https://w463ug1m4g.execute-api.us-east-1.amazonaws.com/prod/file_extraction/source/{aws_details['bucket_name']}"
        for selected_file in selected_files_list:
            if selected_file in small_files:
                text = get_streaming_data_from_aws(selected_file)
                dataFrame.loc[len(dataFrame)] = {"File Name": selected_file, "text": text}
            else:
                large_files.append(selected_file)
        print("Before Streaming...")
        setGetStreamingDataframe.set_stream_dataframe(dataFrame)
        if large_files:
            headers = {"id": aws_details['access_key'], "key": aws_details['secret_key'], "source_paths": ",".join(large_files)}
            response = requests.get(url, headers=headers)
            presigned_urls = json.loads(response.content.decode())
            # fileTest(presigned_urls)
        return JSONResponse({"status": "success", "message": f"{len(selected_files_list)} files successfully uploaded"})
    except Exception as e:
        raise HTTPException(
            status_code=400,
            detail={"status": "error", "message": str(e)}
        )


@app.post("/getting_files_from_azure")
async def GettingFilesFromAzure(
    accountName: str = Form(...),
    accountKey: str = Form(...),
    container: str = Form(...)
):
    try:
        azure_details = {'accountName': accountName, 'accountKey': accountKey, 'container': container}
        setGetAzureDetails.set_account_name(azure_details['accountName'])
        setGetAzureDetails.set_account_key(azure_details['accountKey'])
        setGetAzureDetails.set_container(azure_details['container'])
        blob_service_client = BlobServiceClient(
            account_url=f"https://{azure_details['accountName']}.blob.core.windows.net",
            credential=azure_details['accountKey'])
        container_client = blob_service_client.get_container_client(azure_details['container'])
        blobs_list = [blob.name for blob in container_client.list_blobs()]
        temp_cloud_files.set_files(blobs_list)
        return JSONResponse({"status": "success", "files": blobs_list})
    except Exception as e:
        raise HTTPException(
            status_code=400,
            detail={"status": "error", "message": "Please enter valid details"}
        )


@app.post("/azure_selected_files")
async def AzureSelectedFiles(files: List[str] = Form(...)):
    try:
        # =================================
        currFolder = common.temp_folder
        print("currFolder", currFolder)
        common.delete_folder_files(os.path.join("..",currFolder), True)
        # =================================
        selected_files_list = files
        if not selected_files_list:
            raise HTTPException(
                status_code=400,
                detail={
                    "status": "error",
                    "message": "Please select a file and click upload!",
                    "files": temp_cloud_files.get_files()
                }
            )
        azure_details = setGetAzureDetails.get_()
        blob_service_client = BlobServiceClient(
            account_url=f"https://{azure_details['account_name']}.blob.core.windows.net",
            credential=azure_details['account_key'])
        dataFrame = pd.DataFrame(columns=common.dataframe_columns)
        for file in selected_files_list:
            blob_client = blob_service_client.get_blob_client(container=azure_details['container'], blob=file)
            blob_data = blob_client.download_blob()
            if (blob_data.size <= 150000) and (file.split('.')[-1] in ['txt', 'pdf', 'docx']):
                text = get_streaming_data_from_azure(blob_data, file)
                dataFrame.loc[len(dataFrame)] = {"File Name": file, "text": text}
            else:
                with open(f"./temp_files_to_be_extracted/{file}", "wb") as my_blob:
                    blob_data.readinto(my_blob)
        setGetStreamingDataframe.set_stream_dataframe(dataFrame)
        return JSONResponse({"status": "success", "message": f"{len(selected_files_list)} files successfully uploaded"})
    except Exception as e:
        raise HTTPException(
            status_code=400,
            detail={"status": "error", "message": str(e)}
        )


def get_streaming_data_from_aws(file_name):
    try:
        print("Get Streaming Data from AWS - file_name", file_name)
        aws_details = setGetBucketName.get_()
        s3_client = boto3.client('s3', aws_access_key_id=aws_details['access_key'],
                                 aws_secret_access_key=aws_details['secret_key'])
        s3_object = s3_client.get_object(Bucket=aws_details['bucket_name'],
                                         Key=file_name)
        data = s3_object['Body'].read()
        # print("Data from AWS: ", data)

        # ====================

        local_file_path = f"../temp_files_to_be_extracted/{file_name}"

        with open(local_file_path, 'wb') as f:
            f.write(data)
        print(f"File downloaded and saved at: {local_file_path}")

        # ===================

        if file_name.split('.')[-1] == 'pdf':
            print("PDF Data from AWS...")
            pdf_reader = PyPDF2.PdfReader(BytesIO(data))
            text = ""
            for page in range(len(pdf_reader.pages)):
                page_obj = pdf_reader.pages[page]
                text += page_obj.extract_text()
            print("Extracted Text: ", text)
            return text

        elif file_name.split('.')[-1] == 'docx':
            doc = docx.Document(BytesIO(data))
            para_text = []
            for para in doc.paragraphs:
                para_text.append(para.text)
            doc_text = '\n'.join(para_text)
            print("DOCX from AWS \n Extracted Text :", doc_text)
            return doc_text

        elif file_name.split('.')[-1] == 'txt':
            text = data.decode()
            text = text.replace("\r", " ").replace("\n", " ")
            text = re.sub("\s+", " ", text)
            print("Text from AWS \n Extracted Text: ", text)
            return text
    except Exception as exp:
        print(traceback.format_exc())
        return str(exp)


def get_streaming_data_from_azure(blob_data, file_name):
    try:
        # ====================

        local_file_path = f"../temp_files_to_be_extracted/{file_name}"

        with open(local_file_path, 'wb') as f:
            f.write(blob_data.readall())
        print(f"File downloaded and saved at: {local_file_path}")

        # ===================
        if file_name.split('.')[-1] == 'pdf':
            try: 
                pdf_reader = PyPDF2.PdfReader(BytesIO(blob_data.readall()))
                text = ""
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text()
                return text
            except Exception as e:
                print("File is Empty: ",str(e))
        elif file_name.split('.')[-1] == 'docx':
            doc = docx.Document(BytesIO(blob_data.readall()))
            para_text = []
            for para in doc.paragraphs:
                para_text.append(para.text)
            doc_text = '\n'.join(para_text)
            return doc_text
        elif file_name.split('.')[-1] == 'txt':
            text = blob_data.content_as_text(max_concurrency=1, encoding='UTF-8')
            text = text.replace("\r", " ").replace("\n", " ")
            text = re.sub("\s+", " ", text)
            return text
    except Exception as exp:
        print(traceback.format_exc())
        return str(exp)


from google_drive import GoogleDriveClient
drive_client = None
 
class FolderDetails(BaseModel):
    name: str
    id: str

class GoogleDriveFolder(BaseModel):
    folder_name: FolderDetails
 
@app.get("/connect_google_drive/")
def connect_google_drive():
    """
    Step 1: Generate and return the Google OAuth URL.
    """
    global drive_client
    print("Initiating Google Drive OAuth flow...")
    try:
        drive_client = GoogleDriveClient()
        if drive_client.authenticate():
            print("Authenticated...")
            folders = drive_client.list_files(only_folders = True)
            print("Folder Collected...")
            folders_with_name_and_id = []
            for folder in folders:
                folders_with_name_and_id.append(
                    {
                        "name": folder["name"],
                        'id': folder["id"]
                    }
                )
            return {"message" : "Successfully Authenticated...", "folders": folders_with_name_and_id}
        return {"message": "authentication Failed..."}
    except Exception as e:
        import sys
        print(f"Error adding header: {e}")
        exc_type, exc_obj, exc_tb = sys.exc_info()
        fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
        message = f"{fname} : Line no {exc_tb.tb_lineno} - {exc_type} : {e}"
        raise ValueError(message)
        raise HTTPException(status_code=500, detail=f"Failed to initiate OAuth: {str(e)}")
 
 
@app.post("/google_drive_auth_callback/")
def google_drive_auth_callback(code: str = Query(..., description="Authorization code from Google")):
    """
    Step 2: Accept authorization code and complete the authentication.
    """
    global drive_client
    print("Received OAuth callback with code:", code)
 
    if drive_client is None:
        raise HTTPException(status_code=400, detail="Google Drive client not initialized. Call /connect_google_drive first.")
   
    try:
        success = drive_client.check_authentication(code)
        if success:
            return {"message": "Google Drive authentication successful."}
        else:
            raise HTTPException(status_code=401, detail="Authentication failed.")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"OAuth process failed: {str(e)}")
   
 
@app.post("/getting_files_from_google_drive/")
def get_files_from_google_drive(data: GoogleDriveFolder):
    print("Received folder name:", data.folder_name)
   
    global drive_client
    if drive_client is None:
        connect_google_drive()
 
    folder_name = data.folder_name.name
    folder_id = data.folder_name.id
 
    try:
        result = drive_client.list_files(folder_id)
        return_result = [
            {"name": file["name"], "id": file["id"]}
            for file in result
            if file["name"].lower().endswith((".pdf", ".doc", ".docx", ".jpeg", ".png", ".csv", ".xlsx"))
        ]
        return {
            "message": f"{len(return_result)} Files Names Successfully Retrieved..",
            "files": return_result
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
 
class FileItem(BaseModel):
    name: str
    id: str
 
class FileListRequest(BaseModel):
    files: List[FileItem]  # List of file paths or names
 
@app.post("/streaming_files_from_google_drive/")
def stream_file_from_google_drive(requested_files: FileListRequest):
    print("Received file names:", requested_files.files)
    if drive_client is None:
        print("Drive client is not initialized, initializing now...")
        return {
            "message": "Drive client is not initialized, initializing now..."
        }
 
 
    try:
        currFolder = common.temp_folder
        print("currFolder", currFolder)
        common.delete_folder_files(currFolder, True)
        for file in requested_files.files:
            print(file.name, file.id)
            drive_client.download_file(file_id=file.id, file_name=file.name)
 
        return {
            "message": f"{len(requested_files.files)} files downloaded successfully",
            "files": [file.dict() for file in requested_files.files]
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
 

@app.options("/{path:path}")
async def handle_options_request(path: str):
    return JSONResponse({'status': 'preflight'})


@app.middleware("http")
async def add_cache_control_headers(request: Request, call_next):
    response = await call_next(request)
    response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
    return response

if __name__ == '__main__':
    import uvicorn
    uvicorn.run(app, host='0.0.0.0', port=8001)